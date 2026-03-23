from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / "zust_undergraduate_thesis_template.docx"
HEADER_TEXT = "浙江科技大学计算机科学与技术学院    2026 届本科毕业设计（论文）"


def set_run_font(run, size=12, bold=False, east_asia_font="宋体", ascii_font="Times New Roman"):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    run.font.size = Pt(size)
    run.bold = bold


def configure_document(document):
    section = document.sections[0]
    set_section_page_layout(section)
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(12)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.bold = True
        if style_name == "Heading 1":
            style.font.size = Pt(16)
        elif style_name == "Heading 2":
            style.font.size = Pt(14)
        else:
            style.font.size = Pt(12)


def set_section_page_layout(section):
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def set_section_page_number(section, start=1, fmt="decimal"):
    sect_pr = section._sectPr
    pg_num_types = sect_pr.xpath("./w:pgNumType")
    if pg_num_types:
        pg_num_type = pg_num_types[0]
    else:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))
    pg_num_type.set(qn("w:fmt"), fmt)


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_run_font(run, size=12)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def configure_header_footer(section, show_header_footer=True):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    footer = section.footer
    header.paragraphs[0].clear()
    footer.paragraphs[0].clear()

    if not show_header_footer:
        return

    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_run = header_paragraph.add_run(HEADER_TEXT)
    set_run_font(header_run, size=10, east_asia_font="宋体")

    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    left_run = footer_paragraph.add_run("- ")
    set_run_font(left_run, size=10, east_asia_font="宋体")
    add_page_field(footer_paragraph)
    right_run = footer_paragraph.add_run(" -")
    set_run_font(right_run, size=10, east_asia_font="宋体")


def add_paragraph(document, text="", size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.line_spacing = 1.5
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True, east_asia_font="黑体")
    return paragraph


def add_placeholder_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, bold=False)
    return paragraph


def add_table_of_contents(document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word 中右键目录并选择“更新域”，自动生成目录。"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(placeholder)
    run._r.append(fld_char_end)


def build_cover(document):
    for _ in range(4):
        add_paragraph(document, "")
    add_paragraph(document, "浙江科技大学计算机科学与技术学院", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "本科毕业设计（论文）", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "")
    add_paragraph(document, "论文题目：校园智能问答助手设计与实现", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "")
    add_paragraph(document, "学生姓名：__________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "学    号：__________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "专    业：计算机科学与技术", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "指导教师：__________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "完成日期：____年__月", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    document.add_page_break()


def build_abstract_pages(document):
    add_heading(document, "中文摘要", level=1)
    add_placeholder_paragraph(
        document,
        "【摘要撰写提示】中文摘要一般不超过300字，应概括课题背景、研究内容、实现方法、主要成果与结论，不写公式、图表和参考文献编号。",
    )
    add_placeholder_paragraph(
        document,
        "【摘要占位】本课题面向浙江科技大学校园服务场景，设计并实现了一套基于大模型与本地知识库的校园智能问答助手。系统围绕校园问答中信息分散、检索效率不足和回答自然度不高等问题，完成了知识库构建、查询改写、检索增强、结构化回答与模块化重构等核心功能设计。实验结果表明，该系统能够在图书馆开放时间、学籍管理、奖学金申请等典型校园问题上提供较为准确、自然且具有一定上下文理解能力的回答。该研究为校园服务智能化提供了一种可行实现方案。",
    )
    add_placeholder_paragraph(document, "关键词：校园问答助手；大语言模型；知识库问答；查询改写；模块化设计")
    document.add_page_break()

    add_heading(document, "Abstract", level=1)
    add_placeholder_paragraph(
        document,
        "[Abstract Notes] The English abstract should be consistent with the Chinese abstract in meaning. Keep it concise, independent, and complete.",
    )
    add_placeholder_paragraph(
        document,
        "[Abstract Placeholder] This thesis designs and implements a campus intelligent question-answering assistant based on a large language model and a local knowledge base for campus service scenarios. The system addresses problems such as scattered information, insufficient retrieval efficiency, and weak answer naturalness in campus Q&A. It integrates knowledge base construction, query rewriting, retrieval enhancement, structured response generation, and modular refactoring. Experimental results show that the system can provide relatively accurate, natural, and context-aware answers for typical campus questions such as library opening hours, academic status management, and scholarship applications. This study provides a feasible solution for the intelligent upgrade of campus services.",
    )
    add_placeholder_paragraph(document, "Keywords: campus question answering; large language model; knowledge base QA; query rewriting; modular design")
    document.add_page_break()


def build_toc_page(document):
    add_heading(document, "目录", level=1)
    add_table_of_contents(document)
    document.add_page_break()


def build_body(document):
    chapters = [
        (
            "第1章 引言",
            [
                "1.1 选题背景与研究意义",
                "1.2 国内外研究现状",
                "1.3 研究内容与目标",
                "1.4 论文结构安排",
            ],
        ),
        (
            "第2章 相关技术与理论基础",
            [
                "2.1 大语言模型技术概述",
                "2.2 知识库问答相关技术",
                "2.3 向量数据库与文本切分技术",
                "2.4 Web交互与系统开发框架",
            ],
        ),
        (
            "第3章 需求分析",
            [
                "3.1 系统建设目标",
                "3.2 功能需求分析",
                "3.3 非功能需求分析",
                "3.4 可行性分析",
            ],
        ),
        (
            "第4章 系统总体设计",
            [
                "4.1 系统架构设计",
                "4.2 功能模块设计",
                "4.3 数据流程设计",
                "4.4 数据存储与知识库设计",
            ],
        ),
        (
            "第5章 系统详细设计与实现",
            [
                "5.1 知识库构建模块实现",
                "5.2 查询改写模块实现",
                "5.3 检索增强模块实现",
                "5.4 自然化回答生成模块实现",
                "5.5 系统模块化重构实现",
            ],
        ),
        (
            "第6章 系统测试与结果分析",
            [
                "6.1 测试环境与测试方法",
                "6.2 功能测试",
                "6.3 典型问答效果分析",
                "6.4 系统存在的问题与分析",
            ],
        ),
        (
            "第7章 结束语",
            [
                "7.1 研究工作总结",
                "7.2 创新点与特色",
                "7.3 不足与后续展望",
            ],
        ),
    ]

    for chapter_title, sections in chapters:
        add_heading(document, chapter_title, level=1)
        add_placeholder_paragraph(
            document,
            "【写作提示】本章应围绕章标题展开，段落通顺，重点突出。章标题建议另起一页，标题不要使用标点符号。",
        )
        for section_title in sections:
            add_heading(document, section_title, level=2)
            add_placeholder_paragraph(document, "【内容占位】请在此补充本节正文内容。写作时可根据需要扩展至第三级标题（如 1.1.1），不建议超过第三级标题。")
        document.add_page_break()


def build_acknowledgement(document):
    add_heading(document, "致谢", level=1)
    add_placeholder_paragraph(
        document,
        "【致谢占位】感谢指导教师在课题选题、系统设计、论文撰写与修改过程中给予的悉心指导；感谢学院老师在毕业设计过程中提供的帮助；感谢同学与家人在研究和学习过程中给予的支持与鼓励。",
    )
    document.add_page_break()


def build_references(document):
    add_heading(document, "参考文献", level=1)
    add_placeholder_paragraph(
        document,
        "【参考文献要求】一般不少于10篇，其中外文文献不少于2篇。正文中必须使用上标形式引用文献编号，参考文献格式应符合 GB7714—1987 的书写规范。",
    )
    sample_references = [
        "[1] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[J]. Advances in Neural Information Processing Systems, 2020, 33: 9459-9474.",
        "[2] Zhao W X, Zhou K, Li J, et al. A Survey of Large Language Models[EB/OL]. https://arxiv.org/abs/2303.18223, 2023.",
        "[3] 刘洋, 张伟. 基于大模型的智能问答系统研究[J]. 计算机工程与应用, 2024, 60(8): 15-24.",
        "[4] 陈某某. 智能问答系统设计与实现[D]. 杭州: 某高校, 2023.",
    ]
    for reference in sample_references:
        add_placeholder_paragraph(document, reference)
    document.add_page_break()


def build_appendices(document):
    add_heading(document, "附录1", level=1)
    add_placeholder_paragraph(document, "【附录占位】可放系统关键代码、流程图、数据表、测试用例截图等内容。")
    document.add_page_break()

    add_heading(document, "附录2", level=1)
    add_placeholder_paragraph(document, "【附录占位】如无第二个附录，可删除本页。")


def main():
    document = Document()
    configure_document(document)
    configure_header_footer(document.sections[0], show_header_footer=False)
    build_cover(document)

    prelim_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_page_layout(prelim_section)
    set_section_page_number(prelim_section, start=1, fmt="upperRoman")
    configure_header_footer(prelim_section, show_header_footer=True)
    build_abstract_pages(document)
    build_toc_page(document)

    body_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_page_layout(body_section)
    set_section_page_number(body_section, start=1, fmt="decimal")
    configure_header_footer(body_section, show_header_footer=True)
    build_body(document)
    build_acknowledgement(document)
    build_references(document)
    build_appendices(document)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
